import streamlit as st
import pandas as pd
import openai
import json
from fpdf import FPDF
from docx import Document
import io

# 1. INITIALIZATION & STYLING
st.set_page_config(page_title="EduDiagnostic Elite 2026", layout="wide", page_icon="🎓")

if "OPENAI_API_KEY" not in st.secrets:
    st.error("Missing API Key! Please add 'OPENAI_API_KEY' to Streamlit Secrets.")
    st.stop()

client = openai.OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

# 2. BEAUTY LAYER: FORMATTER (Removes all JSON/Brackets for Printing)
def format_exam_for_humans(json_data, lo_title):
    try:
        data = json_data if isinstance(json_data, dict) else json.loads(json_data)
        paper = f"DIAGNOSTIC ASSESSMENT: {lo_title.upper()}\n"
        paper += "NAME: __________________________    DATE: __________\n"
        paper += "="*60 + "\n\n"
        
        for i, q in enumerate(data.get('questions', []), 1):
            paper += f"Q{i} [{q.get('level', 'Core')}]: {q['question']}\n\n"
            for j, opt in enumerate(q.get('options', [])):
                paper += f"   {chr(65+j)}) {opt}\n"
            paper += "\n" + "-"*30 + "\n\n"
        return paper
    except:
        return "Formatting Error. Please check the raw JSON output."

# 3. DOWNLOAD UTILITIES (PDF & Word)
def get_pdf_bytes(text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=11)
    safe_text = text.encode('latin-1', 'replace').decode('latin-1')
    pdf.multi_cell(0, 8, txt=safe_text)
    return bytes(pdf.output())

def get_docx_bytes(text):
    doc = Document()
    doc.add_heading('Classroom Diagnostic Document', 0)
    doc.add_paragraph(text)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# 4. THE MULTI-AGENT ENGINE
def agent_psychometrician(lo, count, tiers):
    """Creates the tiered assessment with misconception mapping."""
    prompt = f"Create a {count}-question diagnostic for LO: {lo}. Tiers: {tiers}. Output ONLY JSON: 'questions': [{{id, level, question, options:[], correct, misconception_map:{{index:reason}} }}]"
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "system", "content": "You are a Senior Psychometrician. Output clean JSON."},
                  {"role": "user", "content": prompt}],
        response_format={"type": "json_object"}
    )
    return json.loads(response.choices[0].message.content)

def agent_neuro_diagnostician(lo, student_data):
    """Deep analysis: Identifies 'Mental Viruses' and 'Cognitive Conflict' remediation."""
    data_json = student_data.to_json(orient='records')
    prompt = f"""
    ROLE: Elite Educational Neuro-Diagnostician.
    LO: {lo}
    TASK: Analyze these student results and generate a 3-Phase Deep Report.
    PHASE 1: Identify 'Mental Viruses' (Correlate wrong answers to find specific logic flaws).
    PHASE 2: Precision Remediation (Provide 'Cognitive Conflict' activities like the Van Helmont experiment).
    PHASE 3: Individual Risk Profiles (Predict failure points for the NEXT chapter).
    DATA: {data_json}
    """
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "system", "content": "You are a Diagnostic Data Scientist."},
                  {"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content

# 5. USER INTERFACE (TABS)
st.title("🎓 EduDiagnostic Multi-Agent System")
st.markdown("### Better than Standard EI: Diagnosing the 'Why', not just the 'What'.")
tab1, tab2 = st.tabs(["🏗️ Step 1: Create Assessment", "📊 Step 2: Analyze Results"])

# --- TAB 1: CREATE ASSESSMENT ---
with tab1:
    lo_input = st.text_input("Enter Learning Outcome (LO)", "Photosynthesis & Biomass", key="lo1")
    c1, c2 = st.columns(2)
    q_num = c1.slider("Test Length", 5, 15, 7)
    tiers = c2.multiselect("Difficulty Tiers", ["Foundation", "Understanding", "Analytical", "Mastery"], ["Foundation", "Analytical"])
    
    if st.button("🚀 Agent 1: Generate Printable Test"):
        with st.spinner("Architecting questions..."):
            raw_test = agent_psychometrician(lo_input, q_num, tiers)
            st.session_state.clean_test = format_for_print(raw_test, lo_input)
            st.success("Clean Test Paper Ready!")

    if 'clean_test' in st.session_state:
        st.text_area("Preview (No Brackets)", st.session_state.clean_test, height=300)
        c3, c4 = st.columns(2)
        with c3:
            st.download_button("📥 Download PDF Test", data=get_pdf_bytes(st.session_state.clean_test), file_name="Test.pdf")
        with c4:
            st.download_button("📥 Download Word Test", data=get_docx_bytes(st.session_state.clean_test), file_name="Test.docx")

# --- TAB 2: ANALYZE RESULTS ---
with tab2:
    st.subheader("Deep Misconception Analysis")
    uploaded = st.file_uploader("Upload Student Marks (Excel)", type=["xlsx"])
    
    if uploaded:
        df = pd.read_excel(uploaded)
        if st.button("🧠 Agent 2: Run Deep Diagnostic"):
            with st.spinner("Identifying cognitive patterns..."):
                deep_report = agent_neuro_diagnostician(lo_input, df)
                st.session_state.final_report = deep_report
                st.markdown("---")
                st.markdown(deep_report)

    if 'final_report' in st.session_state:
        st.download_button("📥 Download Full Diagnostic Report (PDF)", 
                          data=get_pdf_bytes(st.session_state.final_report), 
                          file_name="Deep_Diagnostic_Report.pdf")
